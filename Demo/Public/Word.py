# Word.py  ——  带格式控制 + 统一返回结构
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn, OxmlElement

from Public.decoration import Decorator   # 你自己的装饰器


class WordDriver:
    """
    统一返回格式：
    {"ret_val": True/False, "data": [...], "info": "normal operation"/异常信息}
    """

    def __init__(self,
                 table_font_name="宋体",
                 table_font_size=10.5,   # 五号字
                 table_align=WD_ALIGN_PARAGRAPH.CENTER,
                 col_widths=None,        # 列表，单位 cm，例 [4, 6]
                 auto_adjust=True):      # 是否让表格自动撑满页面
        # ---- 表格样式 ----
        self.table_font_name = table_font_name
        self.table_font_size = Pt(table_font_size)
        self.table_align = table_align
        self.col_widths = col_widths
        self.auto_adjust = auto_adjust

    # -------------------- 私有：统一单元格格式 --------------------

    def _fmt_cell(self, cell):
        """字体、字号、水平居中、垂直居中、边框"""
        # 1. 字体 & 水平居中
        for paragraph in cell.paragraphs:
            paragraph.alignment = self.table_align   # 水平居中
            for run in paragraph.runs:
                run.font.name = self.table_font_name
                run.font.size = self.table_font_size

        # 2. 垂直居中
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcValign = OxmlElement('w:vAlign')
        tcValign.set(qn('w:val'), 'center')
        tcPr.append(tcValign)

        # 3. 黑色实线边框（0.5 磅）
        self._set_cell_border(
            cell,
            top={"val": "single", "sz": 4, "color": "000000"},
            bottom={"val": "single", "sz": 4, "color": "000000"},
            start={"val": "single", "sz": 4, "color": "000000"},
            end={"val": "single", "sz": 4, "color": "000000"},
        )

    def _apply_width(self, table):
        """应用列宽或自动调整"""
        if self.col_widths:
            for idx, w in enumerate(self.col_widths):
                table.columns[idx].width = Cm(w)
        if self.auto_adjust:
            table.autofit = True
            table.allow_autofit = True

    def _set_cell_border(self, cell, **kwargs):
        """
        通用边框设置函数
        用法:
            _set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": "000000"},
                bottom={"val": "single", "sz": 4, "color": "000000"},
                start={"val": "single", "sz": 4, "color": "000000"},
                end={"val": "single", "sz": 4, "color": "000000"},
            )
        sz = 4  对应 0.5 磅（8 分之一磅单位）
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # 创建 <w:tcBorders>
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
            if edge in kwargs:
                node = OxmlElement(f'w:{edge}')
                for k, v in kwargs[edge].items():
                    node.set(qn(f'w:{k}'), str(v))
                tcBorders.append(node)

        # 清除旧边框再追加
        for child in tcPr.xpath('w:tcBorders'):
            tcPr.remove(child)
        tcPr.append(tcBorders)

    # -------------------- 写纯文本（带标题） --------------------
    @Decorator.raise_err()
    def write_content(self, writer, filename, heading="分子坐标", level=0):
        document = Document()
        document.add_heading(heading, level)
        for file_cont in writer:
            for line in file_cont:
                document.add_paragraph(line)
        document.save(filename)
        return {"ret_val": True, "data": [], "info": "normal operation"}

    # -------------------- 写表格（自动套用格式） --------------------
    @Decorator.raise_err()
    def write_table(self, writer, filename, heading="分子频率", level=0):
        """writer是二维列表，第一行是标题，其余行是数据;
        每行也可以再是二维数组，放在单元格中换行"""
        document = Document()
        # 添加标题
        document.add_heading(heading, level)
        # 添加表格行列
        rows, cols = len(writer), len(writer[0])
        table = document.add_table(rows=rows, cols=cols)

        # 填数据 + 统一格式
        for r_idx, row_data in enumerate(writer):
            row_cells = table.rows[r_idx].cells   # 获取word行对象
            for c_idx, val in enumerate(row_data):  # 获得行的内容
                cell = row_cells[c_idx]
                if r_idx == 0:
                    cell.text = str(val)
                elif isinstance(val, (list, tuple)):
                    # lines = map(str, val)
                    # for line in lines[1:]:
                    #     cell.add_paragraph(line)
                    string_list = [str(item) for item in val]
                    cell.text = "\n".join(string_list)
                else:
                    cell.text = str(val)
                # # 兼容旧接口：若单元格数据本身是列表，换行拼接
                # if isinstance(val, (list, tuple)):
                #     cell.text = "\n".join(map(str, val))
                # else:
                #     cell.text = str(val)
                self._fmt_cell(cell)

        self._apply_width(table)
        document.save(filename)
        return {"ret_val": True, "data": [], "info": "normal operation"}

    # -------------------- 读纯文本 --------------------
    @Decorator.raise_err()
    def read_content(self, filename):
        doc = Document(filename)
        text_list = [p.text for p in doc.paragraphs]
        return {"ret_val": True, "data": ["\n".join(text_list)], "info": "normal operation"}

    # -------------------- 写单段文本 --------------------
    @Decorator.raise_err()
    def write_single_content(self, writer, filename, heading="分子坐标", level=0):
        document = Document()
        document.add_heading(heading, level)
        for file_cont in writer:
            document.add_paragraph(file_cont)
        document.save(filename)
        return {"ret_val": True, "data": [], "info": "normal operation"}


# -------------------------- 自测 --------------------------
if __name__ == '__main__':
    # A = WordDriver(col_widths=[4, 6], auto_adjust=True)  # 自定义列宽（cm）

    # records = [
    #     ["3", '101'],
    #     ["7", '422'],
    #     ["4", '631']
    # ]
    # print(A.write_table(records, "测试表格.docx"))
    # print(A.read_content("SI.docx"))

    wd = WordDriver(
        table_font_name="微软雅黑",
        table_font_size=12,   # 小四
        col_widths=[5, 8],    # 两列分别 5 cm、8 cm
        auto_adjust=True)

    records = [
        ["3", '101'],
        ["7", '422'],
        ["4", '631']
    ]
    print(wd.write_table(records, "测试表格.docx"))
    print(wd.read_content("SI.docx"))
